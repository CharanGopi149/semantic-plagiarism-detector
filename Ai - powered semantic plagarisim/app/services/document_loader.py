from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

import httpx
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from fastapi import HTTPException, UploadFile, status
from pypdf import PdfReader

from app.config import CORPUS_DIR, SUPPORTED_EXTENSIONS
from app.services.text_processing import collapse_whitespace


@dataclass(slots=True)
class SourceDocument:
    name: str
    text: str
    origin: str
    url: str | None = None


def _decode_text(raw_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return raw_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Unable to decode the uploaded text file.")


def _read_pdf(raw_bytes: bytes) -> str:
    reader = PdfReader(BytesIO(raw_bytes))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _read_docx(raw_bytes: bytes) -> str:
    document = DocxDocument(BytesIO(raw_bytes))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def _parse_file_content(filename: str, raw_bytes: bytes) -> str:
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported file type: {extension or 'unknown'}. Use TXT, MD, PDF, or DOCX.",
        )

    if extension in {".txt", ".md"}:
        text = _decode_text(raw_bytes)
    elif extension == ".pdf":
        text = _read_pdf(raw_bytes)
    else:
        text = _read_docx(raw_bytes)

    text = collapse_whitespace(text)
    if not text:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=f"{filename} did not contain readable text.")
    return text


async def load_upload(upload: UploadFile, origin: str) -> SourceDocument:
    if upload.filename is None:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Uploaded file is missing a name.")

    raw_bytes = await upload.read()
    text = _parse_file_content(upload.filename, raw_bytes)
    display_name = Path(upload.filename).stem.replace("_", " ").title()
    return SourceDocument(name=display_name, text=text, origin=origin)


async def fetch_url_source(url: str) -> SourceDocument:
    timeout = httpx.Timeout(25.0, connect=10.0)
    async with httpx.AsyncClient(timeout=timeout, follow_redirects=True) as client:
        response = await client.get(url)
        response.raise_for_status()

    content_type = response.headers.get("content-type", "").lower()
    filename = url.rsplit("/", maxsplit=1)[-1] or "web-source"

    if "application/pdf" in content_type or filename.lower().endswith(".pdf"):
        text = _read_pdf(response.content)
    else:
        soup = BeautifulSoup(response.text, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        text = soup.get_text(separator=" ")

    text = collapse_whitespace(text)
    if not text:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=f"No readable content was found at {url}.")

    title = filename.replace("-", " ").replace("_", " ").title() or "Web Source"
    return SourceDocument(name=title, text=text, origin="url", url=url)


def load_sample_sources() -> list[SourceDocument]:
    sources: list[SourceDocument] = []
    for path in sorted(CORPUS_DIR.glob("*.md")):
        raw_text = path.read_text(encoding="utf-8")
        lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
        if not lines:
            continue
        title = lines[0].replace("#", "").strip() if lines[0].startswith("#") else path.stem.replace("-", " ").title()
        sources.append(SourceDocument(name=title, text=raw_text, origin="sample"))
    return sources


def load_demo_assignment() -> str:
    demo_path = CORPUS_DIR.parent / "demo_assignment.md"
    if not demo_path.exists():
        return ""
    return demo_path.read_text(encoding="utf-8")
